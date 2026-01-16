#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HTML į Word konvertavimo skriptas
Konvertuoja pilnus HTML vertimus į Word dokumentus
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

# Straipsnių informacija
STRAIPSNIAI = [
    {
        "id": 1,
        "html_file": "straipsnis-1-pilnas.html",
        "pavadinimas": "Traumos tyrimai Baltijos šalyse: nuo politinio engimo iki atsigavimo",
        "autoriai": "Evaldas Kazlauskas ir Paulina Zelvienė"
    },
    {
        "id": 2,
        "html_file": "straipsnis-2-pilnas.html",
        "pavadinimas": "Holokausto tarpkartinės pasekmės palikuonių psichinei sveikatai",
        "autoriai": "Patricia Dashorst, Trudy M. Mooren, Rolf J. Kleber, et al."
    },
    {
        "id": 3,
        "html_file": "straipsnis-3-pilnas.html",
        "pavadinimas": "Potrauminis augimas: koncepciniai pagrindai ir empiriniai įrodymai",
        "autoriai": "Richard G. Tedeschi ir Lawrence G. Calhoun"
    },
    {
        "id": 4,
        "html_file": "straipsnis-4-pilnas.html",
        "pavadinimas": "Psichoterapijos metodai PTSS gydymui: kas jiems bendra?",
        "autoriai": "Ulrich Schnyder, Anke Ehlers, Thomas Elbert, et al."
    },
    {
        "id": 5,
        "html_file": "straipsnis-5-pilnas.html",
        "pavadinimas": "Sveikatos priežiūros teikimo traumuotoms populiacijoms iššūkiai",
        "autoriai": "Evaldas Kazlauskas"
    },
    {
        "id": 6,
        "html_file": "straipsnis-6-pilnas.html",
        "pavadinimas": "Psichologinės potrauminio streso sutrikimo teorijos",
        "autoriai": "Chris R. Brewin ir Emily A. Holmes"
    }
]

def extract_content_from_html(html_file):
    """Išgauti turinį iš HTML failo"""
    print(f"Skaitomas HTML failas: {html_file}")

    try:
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # Išgauti pavadinimą
        title_elem = soup.find('h1')
        title = title_elem.get_text() if title_elem else "Nežinomas pavadinimas"

        # Išgauti meta informaciją
        meta_info = {}
        meta_div = soup.find('div', class_='meta-info')
        if meta_div:
            for p in meta_div.find_all('p'):
                text = p.get_text()
                if ':' in text:
                    key, value = text.split(':', 1)
                    meta_info[key.strip()] = value.strip()

        # Išgauti puslapių turinį
        pages = []
        for section in soup.find_all('div', class_='page-section'):
            page_number_elem = section.find('div', class_='page-number')
            page_text_elem = section.find('div', class_='page-text')

            if page_number_elem and page_text_elem:
                page_number = page_number_elem.get_text()
                page_text = page_text_elem.get_text()

                pages.append({
                    'number': page_number,
                    'text': page_text
                })

        return {
            'title': title,
            'meta': meta_info,
            'pages': pages
        }

    except Exception as e:
        print(f"Klaida skaitant HTML failą: {e}")
        return None

def create_word_document(content, output_file):
    """Sukurti Word dokumentą iš HTML turinio"""
    print(f"Kuriamas Word dokumentas: {output_file}")

    try:
        doc = Document()

        # Nustatyti puslapio paraštės
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Pridėti pavadinimą
        title = doc.add_heading(content['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pridėti subpavadinimą
        subtitle = doc.add_heading('PILNAS AUTOMATINIS VERTIMAS (100%)', 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.color.rgb = RGBColor(231, 76, 60)

        doc.add_paragraph()

        # Pridėti meta informaciją
        if content['meta']:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for key, value in content['meta'].items():
                run = meta_para.add_run(f"{key}: ")
                run.bold = True
                meta_para.add_run(f"{value}\n")

        # Pridėti įspėjimą
        warning = doc.add_paragraph()
        warning_run = warning.add_run(
            "⚠ SVARBI PASTABA: Tai yra automatinis vertimas naudojant Google Translate. "
            "Akademiniai terminai ir sudėtingos frazės gali būti išversti ne visai tiksliai. "
            "Svarbiems teiginiams rekomenduojama pasitikrinti su originaliuoniu PDF failu."
        )
        warning_run.font.size = Pt(10)
        warning_run.font.color.rgb = RGBColor(133, 100, 4)

        doc.add_page_break()

        # Pridėti kiekvieno puslapio turinį
        for page in content['pages']:
            # Puslapio numeris
            page_heading = doc.add_heading(page['number'], 2)
            page_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in page_heading.runs:
                run.font.color.rgb = RGBColor(52, 152, 219)

            # Puslapio tekstas
            page_para = doc.add_paragraph(page['text'])
            page_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in page_para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'

            doc.add_paragraph()  # Tarpas tarp puslapių

        # Pridėti footer su informacija
        doc.add_page_break()
        footer_heading = doc.add_heading('Apie šį vertimą', 2)
        footer_para = doc.add_paragraph(
            f"Šis vertimas buvo sukurtas automatiškai naudojant Google Translate API.\n\n"
            f"Iš viso išversta {len(content['pages'])} puslapių akademinio teksto.\n\n"
            f"Vertimas apima visą straipsnio turinį: įvadą, teorinius modelius, "
            f"empirinius tyrimus ir išvadas."
        )

        # Išsaugoti dokumentą
        doc.save(output_file)
        print(f"✓ Word dokumentas sėkmingai sukurtas!")
        return True

    except Exception as e:
        print(f"✗ Klaida kuriant Word dokumentą: {e}")
        return False

def convert_article(straipsnis_info):
    """Konvertuoti vieną straipsnį į Word formatą"""
    html_file = straipsnis_info['html_file']
    output_file = f"straipsnis-{straipsnis_info['id']}-pilnas.docx"

    print("\n" + "="*70)
    print(f" STRAIPSNIS {straipsnis_info['id']}: {straipsnis_info['pavadinimas']}")
    print("="*70)
    print(f"\n📄 HTML: {html_file}")
    print(f"📝 Word: {output_file}")
    print("="*70 + "\n")

    # 1. Išgauti turinį iš HTML
    content = extract_content_from_html(html_file)

    if not content:
        print("✗ Nepavyko išgauti turinio iš HTML!")
        return False

    print(f"✓ Išgauta {len(content['pages'])} puslapių")

    # 2. Sukurti Word dokumentą
    success = create_word_document(content, output_file)

    if success:
        print("\n" + "="*70)
        print(f" ✓✓✓ STRAIPSNIS {straipsnis_info['id']} SĖKMINGAI KONVERTUOTAS! ✓✓✓ ")
        print("="*70)
        print(f"\n📂 Sukurtas failas: {output_file}\n")

    return success

def main():
    """Pagrindinė funkcija - konvertuoti visus straipsnius"""
    print("\n" + "="*70)
    print(" VISŲ STRAIPSNIŲ KONVERTAVIMAS Į WORD FORMATĄ ")
    print("="*70)
    print(f"\nIš viso straipsnių konvertavimui: {len(STRAIPSNIAI)}")
    print("="*70 + "\n")

    success_count = 0
    failed = []

    for straipsnis in STRAIPSNIAI:
        try:
            if convert_article(straipsnis):
                success_count += 1
            else:
                failed.append(straipsnis['id'])
        except Exception as e:
            print(f"✗ KLAIDA konvertuojant straipsnį {straipsnis['id']}: {e}")
            failed.append(straipsnis['id'])

    print("\n" + "="*70)
    print(" ✓✓✓ VISŲ STRAIPSNIŲ KONVERTAVIMAS BAIGTAS! ✓✓✓ ")
    print("="*70)
    print(f"\n📊 Galutinė statistika:")
    print(f"   • Sėkmingai konvertuota: {success_count}/{len(STRAIPSNIAI)}")
    if failed:
        print(f"   • Nepavyko: {', '.join(map(str, failed))}")
    print("\n")

if __name__ == "__main__":
    main()
